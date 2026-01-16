"""
Document Processor Module

Extracts text from PDF and DOCX files for contract analysis.
"""

import io
from typing import Union

from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        file_bytes: Raw bytes of the PDF file
        
    Returns:
        Extracted text content as a string
    """
    reader = PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file_bytes: Raw bytes of the DOCX file
        
    Returns:
        Extracted text content as a string
    """
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n\n".join(text_parts)


def extract_text_from_file(uploaded_file) -> str:
    """
    Extract text from an uploaded file (PDF or DOCX).
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        
    Returns:
        Extracted text content as a string
        
    Raises:
        ValueError: If the file type is not supported
    """
    file_bytes = uploaded_file.read()
    file_name = uploaded_file.name.lower()
    file_type = uploaded_file.type
    
    # Determine file type and extract text
    if file_name.endswith(".pdf") or "pdf" in file_type:
        return extract_text_from_pdf(file_bytes)
    elif file_name.endswith(".docx") or "wordprocessingml" in file_type:
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type: {file_type}. "
            "Please upload a PDF or DOCX file."
        )
